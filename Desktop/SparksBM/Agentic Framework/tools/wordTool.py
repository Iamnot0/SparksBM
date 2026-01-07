"""Word document reading/writing tools"""
import os
from typing import Dict, List, Optional
from docx import Document


class WordTool:
    """Tools for reading and processing Word documents"""
    
    @staticmethod
    def readWord(filePath: str) -> Dict:
        """
        Read Word document and extract text/content
        
        Args:
            filePath: Path to Word file
            
        Returns:
            Dict with text content and metadata
        """
        if not os.path.exists(filePath):
            raise FileNotFoundError(f"Word file not found: {filePath}")
        
        try:
            doc = Document(filePath)
            
            # Extract paragraphs
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            
            # Extract tables
            tables = []
            for table in doc.tables:
                tableData = []
                for row in table.rows:
                    rowData = [cell.text for cell in row.cells]
                    tableData.append(rowData)
                tables.append(tableData)
            
            return {
                'text': '\n'.join(paragraphs),
                'paragraphs': paragraphs,
                'tables': tables,
                'paragraphCount': len(paragraphs),
                'tableCount': len(tables)
            }
        except Exception as e:
            raise RuntimeError(f"Failed to read Word file: {e}")
    
    @staticmethod
    def extractEntities(wordData: Dict, entityType: str = "document") -> List[Dict]:
        """
        Extract entities from Word document for Verinice
        
        Args:
            wordData: Data from readWord
            entityType: Type of entity to extract
            
        Returns:
            List of entity dicts
        """
        entities = []
        
        # Extract from text - look for patterns
        text = wordData.get('text', '')
        
        # Simple extraction - split by sections or paragraphs
        paragraphs = wordData.get('paragraphs', [])
        
        for i, para in enumerate(paragraphs):
            if len(para.strip()) > 20:  # Skip very short paragraphs
                entity = {
                    'type': entityType,
                    'name': f"{entityType}_section_{i}",
                    'description': para[:200],  # First 200 chars
                    'content': para,
                    'rawData': {'paragraphIndex': i}
                }
                entities.append(entity)
        
        # Also extract from tables if present
        for tableIdx, table in enumerate(wordData.get('tables', [])):
            if table and len(table) > 1:  # Has header row
                # Use first row as headers
                headers = table[0]
                for rowIdx, row in enumerate(table[1:], 1):
                    entity = {
                        'type': entityType,
                        'name': f"{entityType}_table_{tableIdx}_row_{rowIdx}",
                        'description': f"Data from table {tableIdx}",
                        'rawData': dict(zip(headers, row))
                    }
                    entities.append(entity)
        
        return entities if entities else [{
            'type': entityType,
            'name': f"{entityType}_document",
            'description': text[:200],
            'content': text
        }]

